import docx,os
from backend_api.actions.translator import translateTo
from backend_api.actions.namegenerator import nameGenerator

doc = docx.Document()
def wordCreater(content,language):
    for line in content:
        line =  translateTo(line,language)[0]
        p = doc.add_paragraph()
        run = p.add_run(line)
        if content[0] == line:
            run.bold = True
            run.font.size = docx.shared.Pt(16)
        else:
            run.font.size = docx.shared.Pt(12)
        run.font.name = 'Arial'
    name = nameGenerator() + '.docx'
    savepath = os.path.join('media','docx',name)
    doc.save(savepath)
    return savepath